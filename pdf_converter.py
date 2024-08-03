import sys
from PyQt5.QtWidgets import QApplication, QWidget, QFileDialog, QMessageBox
from PyQt5.QtGui import QPixmap, QImage
from PyQt5.QtCore import Qt
from pdf_converter_screen import Ui_PDF_Converter
import fitz  # PyMuPDF
from docx import Document
import img2pdf
from PIL import Image

class PdfConverter(QWidget, Ui_PDF_Converter):
    def __init__(self):
        super().__init__()
        self.setupUi(self)
        self.pushButton_select.clicked.connect(self.select_file)
        self.pushButton_convertSave.clicked.connect(self.convert_and_save)
        self.selected_file = None

    def select_file(self):
        options = QFileDialog.Options()
        file, _ = QFileDialog.getOpenFileName(self, "PDF Dosyası Seç", "", "PDF Files (*.pdf);;Word Files (*.docx);;Image Files (*.jpg;*.jpeg;*.png)", options=options)
        if file:
            self.selected_file = file
            self.label_status.setText(f"Seçilen dosya: {file}")

    def convert_and_save(self):
        if not self.selected_file:
            QMessageBox.warning(self, "Uyarı", "Lütfen önce bir dosya seçin.")
            return

        source_format = self.comboBox_source.currentText()
        output_format = self.comboBox_output.currentText()

        if source_format == output_format:
            QMessageBox.warning(self, "Uyarı", "Dönüştürülecek format, dönüşecek format ile aynı olamaz.")
            return

        output_file, _ = QFileDialog.getSaveFileName(self, "Kaydet", "", f"{output_format} Files (*.{output_format.lower()})")
        if not output_file:
            return

        try:
            if source_format == "PDF" and output_format == "WORD":
                self.pdf_to_word(output_file)
            elif source_format == "PDF" and output_format == "JPG":
                self.pdf_to_images(output_file)
            elif source_format == "WORD" and output_format == "PDF":
                self.word_to_pdf(output_file)
            elif source_format == "JPG" and output_format == "PDF":
                self.images_to_pdf(output_file)
            self.label_status.setText(f"Dönüştürülmüş dosya kaydedildi: {output_file}")
        except Exception as e:
            QMessageBox.critical(self, "Hata", f"Dönüştürme işlemi başarısız: {e}")

    def pdf_to_word(self, output_file):
        doc = fitz.open(self.selected_file)
        text = ""
        for page in doc:
            text += page.get_text()
        word_doc = Document()
        word_doc.add_paragraph(text)
        word_doc.save(output_file)

    def pdf_to_images(self, output_file):
        doc = fitz.open(self.selected_file)
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            img.save(f"{output_file[:-4]}_page{page_num + 1}.jpg")

    def word_to_pdf(self, output_file):
        doc = Document(self.selected_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        pdf_doc = fitz.open()
        pdf_doc.new_page(width=595, height=842)  # A4 size in points
        pdf_doc[0].insert_text((72, 72), text)
        pdf_doc.save(output_file)

    def images_to_pdf(self, output_file):
        with open(output_file, "wb") as f:
            f.write(img2pdf.convert([self.selected_file]))

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PdfConverter()
    window.show()
    sys.exit(app.exec_())
